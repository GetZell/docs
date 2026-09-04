from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Zell – Roleplay IT & Compliance Review.docx"

BLUE = "146EF5"
NAVY = "13213C"
INK = "202938"
MUTED = "667085"
LIGHT = "EEF4FF"
PALE = "F7F9FC"
GREEN = "067647"
GREEN_BG = "ECFDF3"
AMBER = "B54708"
AMBER_BG = "FFFAEB"
WHITE = "FFFFFF"
BORDER = "D0D5DD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    run_props.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        run_props.append(u)
    new_run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def style_run(run, size=None, bold=None, color=None, italic=None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")


def add_text(paragraph, text, *, bold=False, color=INK, size=9.3, italic=False):
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color, italic=italic)
    return run


def set_repeat_on_every_page(section) -> None:
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.24)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Heading 1", 18, NAVY, 0, 9),
        ("Heading 2", 12.5, NAVY, 9, 5),
        ("Heading 3", 10, NAVY, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "ZELL", bold=True, color=BLUE, size=9.5)
    add_text(p, "  /  CUSTOMER IT REVIEW  /  ROLEPLAY ONLY", bold=True, color=MUTED, size=7.4)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(7.06))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.7)
    table.columns[1].width = Inches(1.36)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    add_text(left, "CONFIDENTIAL • CURRENT-STATE REVIEW • 3 SEPTEMBER 2026", color=MUTED, size=7.2)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    add_page_number(right)
    for run in right.runs:
        style_run(run, size=7.2, color=MUTED)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, text.upper(), bold=True, color=BLUE, size=8.2)


def add_heading(doc: Document, text: str, level=1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, *, bold_start: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        add_text(p, bold_start, bold=True)
        add_text(p, text[len(bold_start):])
    else:
        add_text(p, text)


def add_bullet(doc: Document, text: str, *, color=INK) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        style_run(run, size=9.2, color=color)
    if not p.runs:
        add_text(p, text, color=color, size=9.2)
    else:
        p.runs[0].text = text


def add_numbered_step(doc: Document, number: int, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.48)
    table.columns[1].width = Inches(6.5)
    badge = table.cell(0, 0)
    set_cell_shading(badge, BLUE)
    set_cell_margins(badge, 75, 90, 75, 90)
    badge.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_text(p, str(number), bold=True, color=WHITE, size=9)
    body = table.cell(0, 1)
    set_cell_margins(body, 55, 145, 55, 60)
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, title + "  ", bold=True, color=NAVY, size=9.5)
    add_text(p, text, color=INK, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc: Document, title: str, text: str, *, kind="info") -> None:
    fill, accent = (LIGHT, BLUE) if kind == "info" else ((GREEN_BG, GREEN) if kind == "success" else (AMBER_BG, AMBER))
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, bold=True, color=accent, size=9.5)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_text(p, text, color=INK, size=9.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_fact_grid(doc: Document, facts: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx in range(0, len(facts), 2):
        row = table.add_row()
        set_cant_split(row)
        for col in range(2):
            cell = row.cells[col]
            set_cell_shading(cell, PALE if ((idx // 2 + col) % 2 == 0) else LIGHT)
            set_cell_margins(cell, 145, 160, 145, 160)
            if idx + col < len(facts):
                label, value = facts[idx + col]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                add_text(p, label.upper(), bold=True, color=MUTED, size=7.4)
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                add_text(p, value, bold=True, color=NAVY, size=10)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_cant_split(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, 105, 105, 105, 105)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, text.upper(), bold=True, color=WHITE, size=7.3)
    for row_idx, row_values in enumerate(rows):
        row = table.add_row()
        set_cant_split(row)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_cell_shading(cell, WHITE if row_idx % 2 == 0 else PALE)
            set_cell_margins(cell, 95, 105, 95, 105)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, bold=(i == 0), color=NAVY if i == 0 else INK, size=7.8)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    doc = Document()
    configure_document(doc)

    # Page 1 — decision brief
    add_kicker(doc, "Security & privacy brief")
    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(5)
    add_text(title, "Roleplay IT &\nCompliance Review", bold=True, color=NAVY, size=30)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    add_text(subtitle, "Voice roleplays only  •  Version 1.0  •  3 September 2026", color=MUTED, size=11)

    add_callout(
        doc,
        "Scope of this review",
        "This document covers only Zell’s browser-based voice roleplay: scenario setup, the simulated conversation, recording and transcription, post-roleplay scoring, feedback, and authorised review. It does not describe or approve any other Zell feature or customer-enabled integration.",
        kind="info",
    )

    add_heading(doc, "Decision summary", level=1)
    add_fact_grid(
        doc,
        [
            ("Primary application", "AWS Milan (eu-south-1)"),
            ("Recording storage", "Google Cloud Storage — EU"),
            ("Voice processing", "ElevenLabs EU-isolated"),
            ("Feedback AI", "AWS Bedrock — EU geography"),
            ("Audio recording", "Stored for review and feedback"),
            ("Customer data training", "Not permitted without written approval"),
            ("Roleplay content", "Synthetic scenario; no prospect call required"),
            ("Contractual framework", "DPA + subprocessor disclosures"),
        ],
    )

    add_heading(doc, "What Zell provides in this scope", level=2)
    add_body(
        doc,
        "Zell provides an AI-supported training environment in which an authorised participant speaks with a simulated persona. Zell records and transcribes the roleplay, evaluates the transcript against the configured scorecard, and returns coaching feedback to authorised users.",
    )
    add_callout(
        doc,
        "GDPR position",
        "Zell operates as the customer’s processor for roleplay personal data. The controls, EU processing configuration, DPA commitments, and subprocessor safeguards described here support a GDPR-compliant deployment. The customer remains responsible for its lawful basis, participant notices, access decisions, and use of results.",
        kind="success",
    )

    # Page 2 — data flow
    add_page_break(doc)
    add_kicker(doc, "01 / Architecture")
    add_heading(doc, "Roleplay data flow", level=1)
    add_numbered_step(doc, 1, "Access", "An authorised participant signs in and selects a customer-configured scenario, persona, and scorecard.")
    add_numbered_step(doc, 2, "Live simulation", "The browser establishes a secure session with the ElevenLabs EU-isolated environment. Participant audio is processed for speech recognition; synthetic voice is returned for the simulated persona.")
    add_numbered_step(doc, 3, "Conversation output", "At the end of the session, the roleplay transcript and session metadata are returned to Zell. Zell retrieves the roleplay recording for customer review.")
    add_numbered_step(doc, 4, "Private storage", "Application and relational data are stored in AWS Milan. Recordings and roleplay files are stored in private Google Cloud Storage buckets configured for the EU.")
    add_numbered_step(doc, 5, "Scoring and feedback", "The transcript, scorecard, and relevant scenario context are processed by EU-hosted AI services to generate scores, explanations, and coaching feedback.")
    add_numbered_step(doc, 6, "Authorised review", "The participant and permitted customer reviewers access the recording, transcript, scores, and feedback through tenant-scoped application controls.")

    add_heading(doc, "Personal data and content in scope", level=2)
    add_matrix(
        doc,
        ["Category", "Examples", "Purpose"],
        [
            ["Account data", "Name, work email, role, tenant/workspace", "Authentication, access and administration"],
            ["Roleplay inputs", "Scenario, persona instructions, scorecard, optional training context", "Configure the simulated conversation and evaluation"],
            ["Voice data", "Live audio and the resulting recording", "Run, replay and review the roleplay"],
            ["Transcript", "Speaker-labelled roleplay dialogue", "Scoring, feedback and authorised review"],
            ["Evaluation data", "Scores, explanations, snippets and coaching feedback", "Training and performance development"],
            ["Technical data", "Session, security, diagnostic and audit metadata", "Operate, protect and troubleshoot the service"],
        ],
        [1.25, 3.15, 2.55],
    )
    add_callout(doc, "Data minimisation", "The roleplay can use a fictional persona and synthetic business scenario. Zell does not require recordings or personal data from real prospects or customers for this use case.", kind="info")

    # Page 3 — processors
    add_page_break(doc)
    add_kicker(doc, "02 / Hosting & subprocessors")
    add_heading(doc, "Roleplay processing locations", level=1)
    add_body(doc, "The following providers are relevant to the roleplay service itself or its supporting security and operational functions.")
    add_matrix(
        doc,
        ["Provider / service", "Role in this use case", "Data involved", "Processing location"],
        [
            ["Amazon Web Services", "Frontend, backend, PostgreSQL database and application infrastructure", "Account, configuration, transcript, scores, feedback, metadata", "Milan, Italy (eu-south-1)"],
            ["AWS Bedrock", "Post-roleplay AI scoring and feedback", "Transcript, scorecard and relevant scenario context", "EU geographic inference profile"],
            ["Google Cloud Storage", "Private object storage for recordings and roleplay files", "Audio recordings and customer-provided files", "European Union multi-region"],
            ["ElevenLabs", "Real-time speech recognition, synthetic voice and conversational orchestration", "Live audio, transcript and roleplay context", "EU-isolated environment"],
            ["Microsoft Azure", "EU-hosted semantic embeddings used in feedback analysis", "Short feedback phrases / derived text", "European Union deployment"],
            ["Cloudflare", "DNS, web delivery and security controls", "IP address, request and security metadata", "Global edge network; nearest data centre"],
            ["Langfuse", "AI quality, prompt and trace monitoring", "Prompt/response content and technical trace metadata", "European Union"],
            ["PostHog", "Product usage analytics", "Usage and device metadata; audio is not intentionally sent", "Frankfurt, Germany / EU"],
            ["Sentry", "Error and performance diagnostics", "Diagnostic and technical metadata; audio is not intentionally sent", "Germany / EU region"],
            ["Resend", "Transactional service email", "Recipient address and message metadata", "Ireland (eu-west-1)"],
        ],
        [1.22, 2.1, 2.15, 1.5],
    )
    add_callout(
        doc,
        "EU processing configuration",
        "Core application data, recordings, voice processing, and AI feedback processing are configured for European locations. Cloudflare’s edge security service is globally distributed; its role is limited to web traffic delivery and protection rather than storage of roleplay recordings.",
        kind="success",
    )

    # Page 4 — security and lifecycle
    add_page_break(doc)
    add_kicker(doc, "03 / Security & lifecycle")
    add_heading(doc, "How roleplay data is protected", level=1)
    add_matrix(
        doc,
        ["Control area", "Current measure"],
        [
            ["Access control", "Authenticated access, role-based permissions, and tenant/workspace scoping restrict customer data to authorised users."],
            ["Transport security", "Roleplay and application traffic is encrypted in transit using TLS-protected provider endpoints."],
            ["Recording storage", "Recordings are held in a private EU storage bucket. Playback access is provided through time-limited signed URLs."],
            ["Secrets and credentials", "Service credentials are kept outside customer-facing application content and access is limited by operational need."],
            ["Monitoring", "Security, application, error, and AI trace monitoring support detection and investigation of abnormal behaviour or failures."],
            ["Isolation", "Application queries and access checks are scoped to the customer tenant; users cannot intentionally browse another tenant’s roleplays."],
            ["Personnel", "Personnel and authorised contractors with access are bound by confidentiality obligations and least-privilege expectations."],
            ["Incident response", "Zell maintains incident handling procedures and will notify the customer without undue delay after becoming aware of a personal-data breach affecting the service."],
        ],
        [1.55, 5.42],
    )

    add_heading(doc, "Retention, deletion and return", level=2)
    add_body(doc, "Is the audio recording retained? ", bold_start="Is the audio recording retained? ")
    add_body(doc, "Yes. Recording is part of the requested roleplay use case and is stored so authorised users can replay and review the exercise.")
    add_body(doc, "How long is data retained? ", bold_start="How long is data retained? ")
    add_body(doc, "Roleplay data is retained for the customer’s contract term or a shorter period agreed with the customer. The customer should define its required operational retention period during implementation.")
    add_body(doc, "What happens when a roleplay is deleted? ", bold_start="What happens when a roleplay is deleted? ")
    add_body(doc, "Deletion removes the application record and related relational data and queues deletion of the corresponding private recording. Provider backups age out through ordinary backup-retention cycles.")
    add_body(doc, "What happens at termination? ", bold_start="What happens at termination? ")
    add_body(doc, "At the customer’s choice and subject to the DPA, Zell returns or deletes personal data after the services end, except where applicable law requires limited retention.")

    # Page 5 — GDPR and AI governance
    add_page_break(doc)
    add_kicker(doc, "04 / Privacy & AI governance")
    add_heading(doc, "Roleplay compliance FAQ", level=1)
    add_heading(doc, "Who is controller and who is processor?", level=2)
    add_body(doc, "The customer is the controller for participant data and decides why the roleplay is used, who participates, and who may review results. Zell acts as processor and processes the data only to provide the contracted roleplay service and on documented instructions.")

    add_heading(doc, "Is customer data used to train general-purpose AI models?", level=2)
    add_body(doc, "No. Zell’s DPA prohibits use of customer personal data to train or improve general-purpose AI models unless the customer gives express written authorisation. Providers are configured and contracted for service processing, not customer-data model training.")

    add_heading(doc, "Does the roleplay make employment decisions?", level=2)
    add_body(doc, "No autonomous employment decision is required or intended. Scores and feedback are training aids. Customers should retain meaningful human review and should not use a roleplay score as the sole basis for a decision with legal or similarly significant effects.")

    add_heading(doc, "How are data-subject requests handled?", level=2)
    add_body(doc, "The customer remains the primary contact for participants. Zell assists with access, correction, deletion, restriction, portability, objection, and other requests where the relevant data is held in the service.")

    add_heading(doc, "Does Zell support DPIAs and audits?", level=2)
    add_body(doc, "Yes. Zell provides information reasonably required for the customer’s data-protection impact assessment and makes compliance information available under the DPA. Audits are handled subject to confidentiality, security, scope, and reasonable-frequency safeguards.")

    add_heading(doc, "How are subprocessors governed?", level=2)
    add_body(doc, "Zell contractually requires subprocessors to protect personal data to standards consistent with Zell’s processor obligations. Zell remains responsible for their performance within the limits of the DPA and provides notice of material changes through its subprocessor disclosure process.")

    add_callout(
        doc,
        "Customer-side requirements",
        "Before launch, the customer should document its lawful basis, give participants an appropriate privacy notice, define retention and reviewer access, avoid unnecessary special-category data, and retain human oversight over the use of scores.",
        kind="info",
    )
    add_heading(doc, "Contact", level=2)
    contact = doc.add_table(rows=1, cols=3)
    contact.alignment = WD_TABLE_ALIGNMENT.CENTER
    contact.autofit = False
    contact.columns[0].width = Inches(0.8)
    contact.columns[1].width = Inches(3.35)
    contact.columns[2].width = Inches(2.8)
    for cell in contact.rows[0].cells:
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, 150, 160, 150, 160)
    avatar = Path(__file__).with_name("avatar.png")
    p = contact.cell(0, 0).paragraphs[0]
    if avatar.exists():
        p.add_run().add_picture(str(avatar), width=Inches(0.55), height=Inches(0.55))
    p = contact.cell(0, 1).paragraphs[0]
    add_text(p, "Moritz Beck\n", bold=True, color=WHITE, size=10.5)
    add_text(p, "CTO & Co-Founder", color="BDCBE0", size=8.6)
    p = contact.cell(0, 2).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_hyperlink(p, "moritz@getzell.com", "mailto:moritz@getzell.com", color="8EB7FF")
    p.add_run("\n")
    add_hyperlink(p, "+49 152 04506322", "tel:+4915204506322", color="8EB7FF")

    doc.core_properties.title = "Zell – Roleplay IT & Compliance Review"
    doc.core_properties.subject = "Customer IT review for the browser-based voice roleplay use case"
    doc.core_properties.author = "Zell UG (haftungsbeschränkt)"
    doc.core_properties.keywords = "Zell, IT review, security, GDPR, roleplay"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
